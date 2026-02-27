#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
특허출원서 DOCX 생성기 (python-docx 사용)
이미지가 제대로 임베딩되도록 python-docx 라이브러리 사용
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# 현재 디렉토리
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DRAWINGS_DIR = os.path.join(BASE_DIR, '도면')

def set_cell_shading(cell, color):
    """셀 배경색 설정"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_styled(doc, text, level=1):
    """한글 스타일 제목 추가"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
    elif level == 2:
        run.font.size = Pt(12)
    else:
        run.font.size = Pt(11)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    return p

def add_paragraph_styled(doc, text, indent=False):
    """본문 단락 추가"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    return p

def create_patent_document():
    doc = Document()

    # 기본 폰트 설정
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # === 표지 ===
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(72)
    run = title.add_run('특 허 출 원 서')
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = '맑은 고딕'

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(24)
    run = subtitle.add_run('(명 세 서)')
    run.font.size = Pt(16)
    run.font.name = '맑은 고딕'

    # 정보 테이블
    doc.add_paragraph()
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    info_data = [
        ('출원인', 'UTTEC'),
        ('발명자', '홍 광선'),
        ('대리인', '(직접 출원)'),
        ('출원일', '2026년 3월 3일'),
        ('문서 버전', '2.1 (하이브리드 AI 추가)')
    ]

    for i, (label, value) in enumerate(info_data):
        cell0 = table.rows[i].cells[0]
        cell1 = table.rows[i].cells[1]
        set_cell_shading(cell0, 'E8E8E8')
        cell0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell0.paragraphs[0].add_run(label)
        run.bold = True
        run.font.name = '맑은 고딕'
        run = cell1.paragraphs[0].add_run(value)
        run.font.name = '맑은 고딕'

    doc.add_page_break()

    # === 발명의 명칭 ===
    add_heading_styled(doc, '【발명의 명칭】')
    p = doc.add_paragraph()
    run = p.add_run('한글: ')
    run.bold = True
    run.font.name = '맑은 고딕'
    run = p.add_run('실시간 콘서트 정보를 시스템 프롬프트로 동적 주입하여 AI 응답을 생성하고 응원봉을 제어하는 스마트 응원봉 시스템 및 방법')
    run.font.name = '맑은 고딕'

    p = doc.add_paragraph()
    run = p.add_run('영문: ')
    run.bold = True
    run.font.name = '맑은 고딕'
    run = p.add_run('Smart Light Stick System and Method for Generating AI Responses and Controlling Light Stick by Dynamically Injecting Real-time Concert Information into System Prompts')
    run.font.name = '맑은 고딕'

    # === 기술분야 ===
    add_heading_styled(doc, '【기술분야】')
    add_paragraph_styled(doc, '본 발명은 스마트 응원봉 기술에 관한 것으로, 더욱 상세하게는 인공지능(AI) 기반 음성 비서 기능을 탑재하여 사용자의 음성 질문에 실시간 콘서트 정보를 기반으로 응답하고, 응답에 포함된 LED 제어 정보를 추출하여 응원봉의 LED를 자동으로 제어하는 시스템 및 방법에 관한 것이다.')

    # === 배경기술 ===
    add_heading_styled(doc, '【발명의 배경이 되는 기술】')
    add_paragraph_styled(doc, '기존 응원봉은 다음과 같은 한계를 가진다:')
    add_paragraph_styled(doc, '1. 일방향 통신: 중앙 제어 시스템에서 응원봉으로 LED 제어 신호만 전송하며, 사용자로부터의 피드백이나 질문을 처리할 수 없다.', indent=True)
    add_paragraph_styled(doc, '2. 정보 제공 불가: "다음 곡이 무엇인지", "현재 곡의 응원색이 무엇인지" 등 콘서트 진행 정보를 사용자에게 능동적으로 제공할 수 없다.', indent=True)
    add_paragraph_styled(doc, '3. 정적 제어 방식: 사전 프로그래밍된 LED 시퀀스만 실행 가능하며, 실시간 상황에 따른 동적 대응이 불가능하다.', indent=True)

    add_heading_styled(doc, '【선행 기술 분석】', level=2)

    # 선행 기술 테이블
    prior_table = doc.add_table(rows=4, cols=3)
    prior_table.style = 'Table Grid'
    headers = ['특허번호', '명칭', '한계점']
    for i, h in enumerate(headers):
        cell = prior_table.rows[0].cells[i]
        set_cell_shading(cell, 'D5E8F0')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = '맑은 고딕'

    prior_data = [
        ('KR102447873B1', '디지털 응원봉 응원 운영시스템', 'AI 기능 없음, 음성 인식 없음'),
        ('KR101822968B1', '응원봉을 구비한 공연 연출 시스템', '중앙 제어만, 양방향 대화 불가'),
        ('US20140184386A1', 'Interactive lighting effect wristband', 'LED 제어만, AI 없음')
    ]
    for i, row_data in enumerate(prior_data, 1):
        for j, cell_data in enumerate(row_data):
            cell = prior_table.rows[i].cells[j]
            run = cell.paragraphs[0].add_run(cell_data)
            run.font.size = Pt(9)
            run.font.name = '맑은 고딕'

    add_paragraph_styled(doc, '상기 선행 기술들은 모두 AI 기반 음성 인터페이스, 실시간 콘텍스트 주입, AI 응답과 LED 제어의 자동 연동 기능이 없다.')

    # === 발명의 내용 ===
    add_heading_styled(doc, '【발명의 내용】')
    add_heading_styled(doc, '【해결하고자 하는 과제】', level=2)
    add_paragraph_styled(doc, '본 발명은 상기와 같은 종래 기술의 문제점을 해결하기 위하여 안출된 것으로, 다음과 같은 과제를 해결하고자 한다:')
    add_paragraph_styled(doc, '1. 사용자가 음성으로 질문하면 현재 콘서트 상황에 맞는 정확한 정보를 AI가 제공하는 시스템 구현', indent=True)
    add_paragraph_styled(doc, '2. AI 응답에서 LED 색상 정보를 자동 추출하여 응원봉을 제어하는 자동화 시스템 구현', indent=True)
    add_paragraph_styled(doc, '3. 콘서트 셋리스트, 아티스트 정보, 현재 곡 인덱스 등을 실시간으로 AI에게 전달하는 동적 콘텍스트 주입 방법 제공', indent=True)

    add_heading_styled(doc, '【과제의 해결 수단】', level=2)
    add_paragraph_styled(doc, '상기 과제를 해결하기 위한 본 발명의 스마트 응원봉 시스템은:')
    add_paragraph_styled(doc, '1. 응원봉 장치: LED, 마이크로컨트롤러, BLE 통신 모듈을 포함', indent=True)
    add_paragraph_styled(doc, '2. 스마트폰 애플리케이션: 음성 인식, AI API 연동, BLE 통신을 처리', indent=True)
    add_paragraph_styled(doc, '3. 로컬 데이터베이스: 콘서트 셋리스트, 아티스트 정보, 현재 곡 인덱스 저장', indent=True)
    add_paragraph_styled(doc, '4. 시스템 프롬프트 생성기: 로컬 데이터를 기반으로 동적 프롬프트 생성', indent=True)
    add_paragraph_styled(doc, '5. AI 서버: 시스템 프롬프트와 사용자 질문을 처리하여 응답 생성', indent=True)

    add_heading_styled(doc, '【발명의 효과】', level=2)
    add_paragraph_styled(doc, '본 발명에 따르면 다음과 같은 효과가 있다:')
    add_paragraph_styled(doc, '1. 실시간 정보 제공, 2. 자연어 인터페이스, 3. LED 자동 제어, 4. 동적 콘텍스트 주입, 5. 빠른 응답 시간(3초 이내)', indent=True)
    add_paragraph_styled(doc, '6. 오프라인 동작 지원, 7. 하이브리드 안정성, 8. 비용 절감, 9. 프라이버시 보호', indent=True)

    doc.add_page_break()

    # === 도면의 간단한 설명 ===
    add_heading_styled(doc, '【도면의 간단한 설명】')
    add_paragraph_styled(doc, '도 1은 본 발명에 따른 AI FanStick 시스템의 전체 구성도이다.')
    add_paragraph_styled(doc, '도 2는 시스템 프롬프트 생성 흐름도이다.')
    add_paragraph_styled(doc, '도 3은 음성-AI-LED 파이프라인 시퀀스 다이어그램이다.')
    add_paragraph_styled(doc, '도 4는 BLE 명령 프로토콜 구조이다.')
    add_paragraph_styled(doc, '도 5는 AI 응답 파싱 알고리즘 흐름도이다.')
    add_paragraph_styled(doc, '도 6은 하이브리드 AI 아키텍처(클라우드 + 온디바이스) 구성도이다.')
    add_paragraph_styled(doc, '도 7은 온디바이스 AI 처리 상세 흐름도이다.')

    # === 구체적인 내용 ===
    add_heading_styled(doc, '【발명을 실시하기 위한 구체적인 내용】')
    add_heading_styled(doc, '1. 시스템 프롬프트 동적 생성 메커니즘', level=2)
    add_paragraph_styled(doc, '본 발명의 시스템 프롬프트 생성부는 로컬 JSON 데이터를 기반으로 현재 콘서트 상태를 반영한 시스템 프롬프트를 동적으로 생성한다. 템플릿: "당신은 {artist_name} 콘서트 AI 비서입니다. 현재 곡: {current_song}, 다음 곡: {next_song}, 답변에 [LED:R,G,B] 포함"')

    add_heading_styled(doc, '2. 음성-AI-LED 통합 파이프라인', level=2)
    add_paragraph_styled(doc, '① 음성 입력 → ② STT 변환(~500ms) → ③ 프롬프트 결합 → ④ AI API 호출(~1500ms) → ⑤ 응답 파싱 → ⑥ TTS 출력 + ⑦ LED 제어(병렬, ~200ms). 전체 처리 시간 약 2.5초.')

    add_heading_styled(doc, '3. AI 응답 LED 색상 추출 알고리즘', level=2)
    add_paragraph_styled(doc, '정규식 패턴: \\[LED:(\\d{1,3}),(\\d{1,3}),(\\d{1,3})\\]. 패턴 매칭 실패 시 로컬 DB에서 현재 곡 응원색 조회하여 폴백 처리.')

    add_heading_styled(doc, '4. 하이브리드 AI 아키텍처', level=2)
    add_paragraph_styled(doc, '(a) 네트워크 상태 감지 → (b) 온라인: 클라우드 AI 우선 → (c) 오프라인: 온디바이스 LLM(Gemma 2B, Q4) 폴백 → (d) 모든 AI 실패: 규칙 기반 템플릿 응답')

    doc.add_page_break()

    # === 청구범위 ===
    add_heading_styled(doc, '【청구범위】')

    claims = [
        ('【청구항 1】 (독립항 - 방법)', '응원봉 장치와 연동된 스마트폰 애플리케이션에서 AI 기반으로 콘서트 정보를 안내하는 방법에 있어서,\n(a) 콘서트 셋리스트, 현재 곡 인덱스, 아티스트 정보, 곡별 응원색 RGB 값을 포함하는 로컬 데이터를 로드하는 단계;\n(b) 상기 로컬 데이터를 기반으로 현재 콘서트 상태를 반영한 시스템 프롬프트를 동적으로 생성하는 단계;\n(c) 사용자의 음성 질문을 텍스트로 변환하는 음성 인식 단계;\n(d) 상기 시스템 프롬프트와 상기 변환된 텍스트를 AI 서버로 전송하여 응답을 수신하는 단계;\n(e) 상기 AI 응답에서 소정의 형식으로 포함된 LED 색상 정보를 추출하는 파싱 단계; 및\n(f) 상기 추출된 LED 색상 정보를 BLE를 통해 응원봉 장치로 전송하여 LED를 제어하는 단계;\n를 포함하는 것을 특징으로 하는 AI 기반 콘서트 정보 안내 방법.'),
        ('【청구항 2】 (종속항)', '제1항에 있어서, 상기 시스템 프롬프트는 LED 색상 응답 형식으로 "[LED:R,G,B]" 형식을 지정하는 지시문을 포함하며, R, G, B는 각각 0-255 범위의 정수값인 것을 특징으로 하는 방법.'),
        ('【청구항 3】 (종속항)', '제1항에 있어서, 상기 (a) 단계의 로컬 데이터는 JSON 형식으로 저장되며, 셋리스트의 각 곡에 대하여 순서, 제목, 응원색 RGB 값, 응원색 이름, 팬 챈트 정보를 포함하는 것을 특징으로 하는 방법.'),
        ('【청구항 4】 (종속항)', '제1항에 있어서, 상기 (c) 단계부터 상기 (f) 단계까지의 전체 처리 시간이 3초 이내인 것을 특징으로 하는 방법.'),
        ('【청구항 5】 (종속항)', '제1항에 있어서, 상기 (e) 단계에서 LED 색상 정보가 추출되지 않는 경우, 현재 곡의 응원색을 유지하거나 로컬 데이터에서 해당 곡의 응원색을 조회하여 적용하는 것을 특징으로 하는 방법.'),
        ('【청구항 6】 (독립항 - 시스템)', '스마트폰 애플리케이션과 응원봉 장치를 포함하는 AI 기반 스마트 응원봉 시스템에 있어서, 상기 스마트폰 애플리케이션은: (a) 로컬 데이터 저장부; (b) 프롬프트 생성부; (c) 음성 인식부; (d) AI 통신부; (e) 응답 파싱부; 및 (f) BLE 통신부;를 포함하고, 상기 응원봉 장치는: (g) 발광부; (h) BLE 수신부; 및 (i) 마이크로컨트롤러;를 포함하는 것을 특징으로 하는 AI 기반 스마트 응원봉 시스템.'),
        ('【청구항 7】 (종속항)', '제6항에 있어서, 상기 프롬프트 생성부는 현재 곡 인덱스가 변경될 때마다 시스템 프롬프트를 재생성하여, AI 서버가 항상 최신 콘서트 상태를 인지할 수 있도록 하는 것을 특징으로 하는 시스템.'),
        ('【청구항 8】 (종속항)', '제6항에 있어서, 상기 BLE 통신부는 "C:R,G,B"(색상), "P:패턴명"(패턴), "T:텍스트"(표시) 형식의 구조화된 명령 프로토콜을 사용하는 것을 특징으로 하는 시스템.'),
        ('【청구항 9】 (독립항 - 장치)', 'BLE 통신 모듈, 복수의 LED, 마이크로컨트롤러를 포함하는 스마트 응원봉 장치에 있어서, 상기 마이크로컨트롤러는: (a) AI 서버 응답에서 추출된 색상 정보 기반 LED 제어 명령 수신; (b) 명령 파싱하여 RGB 값 추출; (c) LED 색상 제어;를 수행하는 것을 특징으로 하는 스마트 응원봉 장치.'),
        ('【청구항 10】 (종속항)', '제9항에 있어서, 상기 마이크로컨트롤러는 rainbow, pulse, blink, wave 패턴을 포함하는 복수의 LED 애니메이션 패턴을 저장하고 실행하는 것을 특징으로 하는 장치.'),
        ('【청구항 11】 (독립항 - 하이브리드 AI)', '클라우드 AI 서버와 온디바이스 AI를 병용하여 응원봉을 제어하는 방법에 있어서, (a) 네트워크 상태 감지; (b) 온라인 시 클라우드 AI 호출; (c) 오프라인 시 온디바이스 경량 LLM(2B 이하 파라미터, 양자화) 호출; (d) 응답에서 LED 제어 정보 추출;을 포함하는 것을 특징으로 하는 하이브리드 AI 기반 스마트 응원봉 제어 방법.'),
        ('【청구항 12】 (종속항)', '제11항에 있어서, 온디바이스 AI도 실패 시 규칙 기반 응답을 생성하는 폴백 단계를 더 포함하는 것을 특징으로 하는 방법.'),
        ('【청구항 13】 (종속항)', '제11항에 있어서, 상기 온디바이스 경량 언어모델은 앱 최초 실행 시 원격 서버로부터 다운로드되어 로컬 저장소에 저장되며, 버전 관리를 통해 업데이트가 가능한 것을 특징으로 하는 방법.'),
        ('【청구항 14】 (종속항)', '제11항에 있어서, 상기 온디바이스 경량 언어모델은 콘서트 도메인에 특화된 데이터셋으로 파인튜닝되어 답변 품질이 향상된 것을 특징으로 하는 방법.')
    ]

    for title, content in claims:
        add_heading_styled(doc, title, level=3)
        add_paragraph_styled(doc, content)

    doc.add_page_break()

    # === 요약서 ===
    add_heading_styled(doc, '【요약서】')
    add_heading_styled(doc, '【요약】', level=2)
    add_paragraph_styled(doc, '본 발명은 AI 기반 스마트 응원봉 시스템 및 방법에 관한 것으로, 콘서트 셋리스트, 현재 곡 인덱스, 아티스트 정보를 포함하는 로컬 데이터를 기반으로 시스템 프롬프트를 동적으로 생성하고, 사용자의 음성 질문과 함께 AI 서버로 전송하여 응답을 수신한다. AI 응답에서 [LED:R,G,B] 형식의 색상 정보를 추출하여 BLE를 통해 응원봉의 LED를 자동 제어한다. 클라우드 AI와 온디바이스 경량 언어모델을 병용하는 하이브리드 아키텍처를 채택하여, 오프라인 환경에서도 콘서트 정보 안내 및 LED 제어가 가능하다.')

    add_heading_styled(doc, '【대표도】', level=2)
    add_paragraph_styled(doc, '도 1')

    doc.add_page_break()

    # === 도면 ===
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('【도 면】')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = '맑은 고딕'

    # 도면 이미지 추가
    drawings = [
        ('도1_전체시스템구성도.png', '【도 1】 AI FanStick 전체 시스템 구성도'),
        ('도2_시스템프롬프트생성흐름도.png', '【도 2】 시스템 프롬프트 생성 흐름도'),
        ('도3_음성AI_LED파이프라인시퀀스.png', '【도 3】 음성-AI-LED 파이프라인 시퀀스 다이어그램'),
        ('도4_BLE명령프로토콜구조.png', '【도 4】 BLE 명령 프로토콜 구조'),
        ('도5_AI응답파싱알고리즘흐름도.png', '【도 5】 AI 응답 파싱 알고리즘 흐름도'),
        ('도6_하이브리드AI아키텍처.png', '【도 6】 하이브리드 AI 아키텍처 (클라우드 + 온디바이스)'),
        ('도7_온디바이스AI처리상세흐름도.png', '【도 7】 온디바이스 AI 처리 상세 흐름도')
    ]

    for filename, caption in drawings:
        doc.add_page_break()

        # 도면 제목
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(caption)
        run.bold = True
        run.font.size = Pt(13)
        run.font.name = '맑은 고딕'

        # 도면 이미지 - 직접 문서에 추가
        img_path = os.path.join(DRAWINGS_DIR, filename)
        if os.path.exists(img_path):
            # 이미지를 직접 문서에 추가 (paragraph 없이)
            doc.add_picture(img_path, width=Inches(6.0))
            # 마지막 paragraph를 가운데 정렬
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            print(f'이미지 추가: {filename}')
        else:
            print(f'이미지 없음: {img_path}')

    # 마무리
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    run = p.add_run('- 끝 -')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = '맑은 고딕'

    # 저장
    output_path = os.path.join(BASE_DIR, 'AI_FanStick_특허출원서_최종.docx')
    doc.save(output_path)
    print(f'\nDOCX 파일 생성 완료: {output_path}')
    print(f'파일 크기: {os.path.getsize(output_path):,} bytes')

if __name__ == '__main__':
    create_patent_document()
