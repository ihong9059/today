"""강사카드 docx 생성 - 홍광선 / 용인 중소기업진흥원 AI 공장자동화 교육"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_korean_font(run, name="맑은 고딕", size=10):
    run.font.name = name
    run.font.size = Pt(size)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:cs"), name)


def set_cell_bg(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_kv_table(doc, rows, label_width=Cm(4.5), value_width=Cm(12)):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(rows):
        c0 = table.cell(i, 0)
        c1 = table.cell(i, 1)
        c0.width = label_width
        c1.width = value_width
        set_cell_bg(c0, "F2F2F2")
        c0.text = ""
        c1.text = ""
        r0 = c0.paragraphs[0].add_run(k)
        r0.bold = True
        set_korean_font(r0, size=10)
        r1 = c1.paragraphs[0].add_run(v)
        set_korean_font(r1, size=10)
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        set_korean_font(run, size=13)
    else:
        set_korean_font(run, size=11)
    return p


def add_para(doc, text, bold=False, size=10, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    set_korean_font(run, size=size)
    return p


def add_bullet(doc, text, size=10):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_korean_font(run, size=size)
    return p


def add_grid_table(doc, header, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(header):
        c = table.cell(0, j)
        set_cell_bg(c, "DDEBF7")
        c.text = ""
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        set_korean_font(run, size=10)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if col_widths:
            c.width = col_widths[j]
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            c = table.cell(i, j)
            c.text = ""
            run = c.paragraphs[0].add_run(val)
            set_korean_font(run, size=10)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if col_widths:
                c.width = col_widths[j]
    return table


def main():
    doc = Document()

    # 페이지 여백
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # 기본 스타일 한글 폰트
    normal = doc.styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal.font.size = Pt(10)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    rFonts.set(qn("w:ascii"), "맑은 고딕")
    rFonts.set(qn("w:hAnsi"), "맑은 고딕")

    # ── 제목
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("강  사  카  드")
    run.bold = True
    set_korean_font(run, size=22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("용인 중소기업진흥원 「AI 공장자동화 프로그래밍」 대면교육 (5회 15시간)")
    run.bold = True
    set_korean_font(run, size=11)

    # ── 1. 기본 인적사항
    add_heading(doc, "1. 기본 인적사항")
    add_kv_table(doc, [
        ("성명 (국문)", "홍 광 선"),
        ("성명 (영문)", "HONG KWANGSUN"),
        ("생년월일", "1960. 09. 15."),
        ("휴대전화", "010-2401-9059"),
        ("이메일", "ihong9059@gmail.com"),
        ("자택 주소", "경기도 수원시 영통구 효원로 363 위브하늘채 아파트 133동 502호"),
        ("소속", "㈜유티텍 (UTTEC Co., Ltd.)"),
        ("직책", "대표이사 / 기업부설연구소 소장"),
        ("사업장 주소", "경기도 용인시 기흥구 흥덕 유타워 2404호"),
        ("사업자등록번호", "684-86-00438"),
    ])

    # ── 2. 학력
    add_heading(doc, "2. 학력")
    add_grid_table(doc,
        header=["기간", "학교 / 학과", "학위"],
        rows=[
            ("", "아주대학교 전자공학과", "석사"),
        ],
        col_widths=[Cm(4), Cm(8), Cm(4.5)],
    )

    # ── 3. 주요 경력
    add_heading(doc, "3. 주요 경력")
    add_grid_table(doc,
        header=["기간", "소속 / 직위", "주요 업무"],
        rows=[
            ("1985 ~ 2003", "삼성전자",
             "임베디드 시스템 개발 (18년)"),
            ("2011 ~ 2015", "유티솔",
             "임베디드 시스템 개발"),
            ("2016 ~ 현재", "㈜유티텍 대표이사 / 기업부설연구소 소장",
             "Network Solution (BLE Mesh · LoRa · LoRa+AI 무선 제어) 전문기업 운영 및 R&D 총괄"),
        ],
        col_widths=[Cm(3.5), Cm(5.5), Cm(7.5)],
    )

    # ── 4. 주요 사업 실적
    add_heading(doc, "4. 주요 사업 실적 (㈜유티텍)")

    add_heading(doc, "4-1. 양산 운영 중인 제품 (5종)", level=2)
    for item in [
        "STM32F756 컴프레서 밸브 컨트롤러",
        "STM32F407 세탁기 컨트롤러",
        "Raspberry Pi CM4 EtherCAT 컨트롤러",
        "Raspberry Pi 3 V-Cut 컨트롤러",
        "nRF52832 BLE 온도 컨트롤러",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4-2. 국내 대형 납품 실적", level=2)
    for item in [
        "현대건설 / 롯데건설 / 한화건설 / LH / 홈플러스 — EMS IoT 시스템 공급",
        "하나금융캠퍼스 — BLE Mesh 조명 제어",
        "필로스 GC / 광릉 CC — 골프장 LoRa 조명 제어",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4-3. 해외 수출 실적", level=2)
    for item in [
        "2017 — 스위스 LUMITEC · 일본 OPTILED 수출",
        "2022 — 일본 하네다공항 아나모리나역 자전거주차장 500대 수출",
        "2023 — 일본 나고야 사카에역 자전거주차장 3,300대 수출",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4-4. 국내 최초 상용화", level=2)
    add_bullet(doc, "2018 — LoRa 콘트롤러 + 클라우드 통합 시스템 국내 최초 상용화")

    # ── 5. 인증 / 특허
    add_heading(doc, "5. 인증 / 특허")
    add_kv_table(doc, [
        ("국내 인증", "한국 KC 인증 (RS485 Modbus RTU 프로토콜 포함)"),
        ("해외 인증", "일본 TELEC 인증 · 유럽 CE 인증"),
        ("특허", "한국 · 일본 특허 2건 등록"),
        ("기업 인증", "기업부설연구소 보유 · 벤처기업 인증 (2018)"),
    ])

    # ── 6. 보유 기술 / 강의 분야
    add_heading(doc, "6. 보유 기술 / 강의 분야")

    add_heading(doc, "6-1. 임베디드 / 하드웨어 (전문)", level=2)
    for item in [
        "Raspberry Pi 4B / CM4 / Zero — GPIO · I2C · SPI · UART 실전",
        "ESP32 (WROOM · C3 · C6) — Wi-Fi · BLE · OTA",
        "STM32 F4 / F7 / H7 — 양산 컨트롤러",
        "nRF52832 / NimBLE — BLE 양산 펌웨어",
        "산업 통신 : RS485 · Modbus RTU · CAN · EtherCAT · LoRa · MQTT",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "6-2. AI / 데이터", level=2)
    for item in [
        "Python · Flask · FastAPI 풀스택",
        "Google Colab — CNN · MNIST · 시계열 분석 노트북 13개+",
        "EasyOCR 번호판 인식 (인식률 90%)",
        "MobileNetV3 볼트 품질검사 AI (재현율 100%)",
        "Claude · GPT 등 LLM API 활용 (이상탐지 보고서 자동 생성 외)",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "6-3. 교육 콘텐츠 제작", level=2)
    for item in [
        "UTTEC Edu 온라인 교육 플랫폼 운영 (37개 코스 · 765일 커리큘럼)",
        "사전빌드 펌웨어 272개 항목 관리 (Arduino-CLI 자동화)",
        "Python Vibe 교육 앱 (AI 코드 생성 → 실행 → 설명)",
        "Remotion 기반 AI 교육 비디오 30편 이상 제작 (한 · 영)",
        "CUDA 6주 커리큘럼 (Jetson Nano · 16개 예제)",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "6-4. 강의 가능 과목", level=2)
    for i, item in enumerate([
        "Raspberry Pi + Python 공장자동화 프로그래밍 ★ (본 교육 과정)",
        "ESP32 / Arduino IoT 펌웨어 개발",
        "Python 기초 ~ 데이터 분석 · 시각화",
        "Google Colab AI 입문 (CNN · 이상탐지)",
        "Claude / GPT API 활용 실무 자동화",
        "산업 통신 (RS485 · Modbus · LoRa) 실습",
    ], start=1):
        add_para(doc, f"  {i}. {item}")

    # ── 7. 강의 / 교육 경력
    add_heading(doc, "7. 강의 / 교육 경력")
    add_grid_table(doc,
        header=["시기", "교육 대상 / 주관", "과목", "형태"],
        rows=[
            ("2025 ~ 현재", "UTTEC Edu (자체 온라인 플랫폼)",
             "C / Python / ESP32 / RPi 통합 커리큘럼", "온라인 (37코스)"),
            ("2026-03", "사내 / 협력사",
             "CUDA 6주 커리큘럼 (Jetson Nano · 16예제)", "자체 자료"),
            ("2026 (예정)", "용인 중소기업진흥원",
             "AI 공장자동화 프로그래밍 (5회 15시간)", "대면 실습"),
        ],
        col_widths=[Cm(2.5), Cm(5.0), Cm(6.5), Cm(3.0)],
    )

    # ── 8. 본 교육 강의 계획 요약
    add_heading(doc, "8. 본 교육 (용인 중소기업진흥원) 강의 계획 요약")
    add_kv_table(doc, [
        ("교육명", "AI 공장자동화 프로그래밍 (Raspberry Pi + Python)"),
        ("대상", "중소기업 종업원 (프로그래밍 초보자)"),
        ("형태", "대면 실습 교육"),
        ("일정", "3시간 × 5회 = 총 15시간"),
        ("인원", "15 ~ 20명"),
        ("사용 하드웨어", "Raspberry Pi 4B + UTTEC Shield Board (자체 설계 PCB)"),
        ("사용 SW", "Python 3.11 · Flask · Google Colab · Claude API"),
        ("1회차", "환경 구축 + Python 기초 (Pi 부팅 · Linux · VNC · Thonny)"),
        ("2회차", "GPIO + 센서 (LED 신호등 · AHT20 온습도 · OLED · WS2812)"),
        ("3회차", "데이터 + 통신 (CSV · matplotlib · ESP32C3 UART · Flask 대시보드)"),
        ("4회차", "AI 적용 (Colab MNIST · Z-score 이상탐지 · Claude API 보고서)"),
        ("5회차", "종합 프로젝트 (조별 미니 스마트팩토리 구축 · 시연 발표)"),
    ])

    add_heading(doc, "본 교육의 차별점", level=2)
    for i, item in enumerate([
        "자체 설계 UTTEC Shield 보드 — 별도 배선 없이 즉시 실습",
        "양산 경험 기반 — 임베디드 경력 38년 · 양산 제품 5종 운영 노하우 전달",
        "현장 적용 가능 — 수강생이 키트 가져가 자기 공장에 바로 프로토타이핑",
    ], start=1):
        add_para(doc, f"  {i}. {item}")

    # ── 9. 자격 / 수상
    add_heading(doc, "9. 자격 / 수상")
    add_para(doc, "  해당 없음")

    # ── 10. 첨부 서류
    add_heading(doc, "10. 첨부 서류")
    for item in [
        "사업자등록증 사본 (㈜유티텍)",
        "기업부설연구소 인정서",
        "벤처기업확인서",
        "주요 특허증 사본 (한국 · 일본)",
        "KC / TELEC / CE 인증서 사본",
        "회사소개서 (UTTEC)",
        "AI 공장자동화 교육 제안서",
    ]:
        add_bullet(doc, item)

    # ── 서약
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "본인은 위 내용이 사실임을 확인하며,\n"
        "용인 중소기업진흥원 「AI 공장자동화 프로그래밍」 교육의 강사로 성실히 임할 것을 서약합니다."
    )
    set_korean_font(run, size=11)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("2026 년       월       일")
    set_korean_font(run, size=11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("강사 :    홍   광   선     (인)")
    run.bold = True
    set_korean_font(run, size=12)

    out_path = r"C:\Users\lenovo\Downloads\강사카드_홍광선.docx"
    doc.save(out_path)
    print(f"SAVED: {out_path}")


if __name__ == "__main__":
    main()
